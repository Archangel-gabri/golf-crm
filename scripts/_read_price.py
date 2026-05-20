import os, shutil, sys
d = r'C:/Users/VADIM/Desktop/'
print("All docx:")
for f in os.listdir(d):
    if f.endswith('.docx'):
        print("  ", repr(f))

src = None
for f in os.listdir(d):
    if f.endswith('.docx') and not f.startswith('~') and ('Крыл' in f or 'Прайс' in f or 'ол' in f):
        src = os.path.join(d, f)
        print("MATCH:", repr(f))
        break
if not src:
    print("NOT FOUND"); sys.exit(1)
dst = r'Z:/_price.docx'
shutil.copy(src, dst)
print("COPIED OK")

from docx import Document
doc = Document(dst)
for p in doc.paragraphs:
    if p.text.strip():
        print("P:", p.text)
for i, t in enumerate(doc.tables):
    print(f"\n=== TABLE {i} ===")
    for row in t.rows:
        print(' | '.join(c.text.strip() for c in row.cells))
